from __future__ import annotations

import platform
import re
import shutil
import subprocess
import tempfile
import time
from contextlib import contextmanager
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

try:
    import pywintypes
except Exception:
    pywintypes = None

try:
    import pythoncom
except Exception:
    pythoncom = None

try:
    import win32com.client as win32
except Exception:
    win32 = None

from docx import Document
from docx.shared import Pt


WORD_SAVE_AS_TEXT = 7
SEPARATOR_RE = re.compile(r"={4,}")
BLOCK_SEPARATOR_RE = re.compile(r"\+{4,}")
ENCODING_CANDIDATES = ("cp1254", "cp1251", "utf-8-sig", "utf-8", "cp1252", "latin1")


@dataclass
class Question:
    prompt: str
    answers: list[str]


@dataclass
class ConversionResult:
    source_name: str
    profile: str
    encoding: str
    backend: str
    question_count: int
    hemis_text: str
    txt_path: Path
    docx_path: Path


OverrideHandler = Callable[[list[str]], list[list[str]]]


def normalize_text(value: str) -> str:
    value = value.replace("\ufeff", "").replace("\u200b", "")
    value = value.replace("\u000b", "\n").replace("\u000c", "\n")
    value = value.replace("\r", "\n")
    value = value.replace("\u00a0", " ")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n+", "\n", value)
    return value.strip()


def normalize_segment(value: str) -> str:
    return re.sub(r"\s+", " ", normalize_text(value)).strip()


def clean_raw_line(value: str) -> str:
    value = value.replace("\ufeff", "").replace("\u200b", "")
    value = re.sub(r"^(?:ï»¿|п»ї)+", "", value)
    return value.strip()


def is_separator_line(value: str, pattern: re.Pattern[str]) -> bool:
    cleaned = clean_raw_line(value)
    if pattern.fullmatch(cleaned):
        return True

    if pattern is BLOCK_SEPARATOR_RE and re.fullmatch(r"[^\w\d]{1,3}\+{4,}", cleaned):
        return True
    if pattern is SEPARATOR_RE and re.fullmatch(r"[^\w\d]{1,3}={4,}", cleaned):
        return True
    return False


def strip_question_number(prompt: str) -> str:
    prompt = normalize_segment(prompt)
    prompt = re.sub(r"^(\d+)\s+\1\s*[\.\)]\s*", "", prompt)
    return re.sub(r"^\d+\s*[\.\)]\s*", "", prompt).strip()


def normalize_answer(answer: str) -> str:
    answer = normalize_segment(answer)
    if answer.startswith("#"):
        return "#" + answer.lstrip("#").strip()
    return answer


def split_blocks(text: str) -> list[list[str]]:
    lines = [clean_raw_line(line) for line in text.splitlines()]
    lines = [line for line in lines if line]
    blocks: list[list[str]] = []
    current: list[str] = []

    for line in lines:
        if is_separator_line(line, BLOCK_SEPARATOR_RE):
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(line)

    if current:
        blocks.append(current)

    return blocks


def split_raw_blocks(text: str) -> list[list[str]]:
    blocks: list[list[str]] = []
    current: list[str] = []

    for raw_line in text.splitlines():
        line = clean_raw_line(raw_line)
        if not line:
            continue
        if is_separator_line(line, BLOCK_SEPARATOR_RE):
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(line)

    if current:
        blocks.append(current)

    return blocks


def block_to_segments(lines: list[str]) -> list[str]:
    segments: list[str] = []
    current: list[str] = []

    for line in lines:
        if is_separator_line(line, SEPARATOR_RE):
            segments.append(normalize_segment(" ".join(current)))
            current = []
            continue
        current.append(line)

    segments.append(normalize_segment(" ".join(current)))
    return [segment for segment in segments if segment]


def split_first_segment_if_needed(segments: list[str]) -> list[str]:
    if segments and "#" in segments[0] and not segments[0].startswith("#"):
        question, answer = segments[0].split("#", 1)
        return [question.strip(), "#" + answer.strip(), *segments[1:]]
    return segments


def split_merged_segment(segment: str, pattern: str) -> tuple[str, str]:
    match = re.search(pattern, segment)
    if not match:
        raise ValueError(f"Split pattern not found in merged segment: {segment}")
    split_at = match.start()
    return normalize_segment(segment[:split_at]), normalize_segment(segment[split_at:])


def chunk_segments(segments: list[str], chunk_size: int = 5) -> list[list[str]]:
    if len(segments) % chunk_size != 0:
        raise ValueError(f"Cannot chunk {len(segments)} segments into {chunk_size}: {segments}")
    return [segments[index:index + chunk_size] for index in range(0, len(segments), chunk_size)]


def make_question(segments: list[str]) -> Question:
    if len(segments) != 5:
        raise ValueError(f"Expected 5 segments, got {len(segments)}: {segments}")

    prompt = strip_question_number(segments[0])
    answers = [normalize_answer(answer) for answer in segments[1:]]

    correct_answers = sum(answer.startswith("#") for answer in answers)
    if correct_answers == 0:
        answers[0] = "#" + answers[0].lstrip("#").strip()
        correct_answers = 1
    if correct_answers != 1:
        raise ValueError(f"Expected exactly one correct answer marker: {segments}")

    return Question(prompt=prompt, answers=answers)


def is_header_line(line: str) -> bool:
    stripped = normalize_segment(line)
    if not stripped or stripped.startswith("#"):
        return False

    lowered = stripped.lower()
    if (
        "plaintext" in lowered
        or "копировать код" in lowered
        or "êîï" in lowered
        or "mavzu:" in lowered
    ):
        return True

    words = stripped.split()
    letters = [char for char in stripped if char.isalpha()]
    if not letters:
        return False

    upper_count = sum(char.isupper() for char in letters)
    lower_count = sum(char.islower() for char in letters)
    return len(words) >= 2 and upper_count >= 3 and lower_count == 0 and not re.match(r"^[A-D]\)", stripped)


def remove_duplicate_number_token(tokens: list[str]) -> list[str]:
    if len(tokens) >= 2 and re.fullmatch(r"\d+", tokens[0]) and re.match(rf"^{tokens[0]}[\.\)]", tokens[1]):
        return tokens[1:]
    return tokens


def uz_override_100(segments: list[str]) -> list[list[str]]:
    answer4, next_question = split_merged_segment(segments[4], r"(?=\d+\.)")
    return [
        [segments[0], segments[1], segments[2], segments[3], answer4],
        [next_question, segments[5], segments[6], segments[7], segments[8]],
    ]


def uz_override_157(segments: list[str]) -> list[list[str]]:
    answer4, next_question = split_merged_segment(segments[4], r"(?=Futbol to'pining og'irligi qancha\?)")
    return [
        [segments[0], segments[1], segments[2], segments[3], answer4],
        [next_question, segments[5], segments[6], segments[7], segments[8]],
    ]


def uz_override_361(segments: list[str]) -> list[list[str]]:
    return chunk_segments(segments)


def uz_override_386(segments: list[str]) -> list[list[str]]:
    return [[segments[0], f"{segments[1]} {segments[2]}", segments[3], segments[4], segments[5]]]


def uz_override_463(segments: list[str]) -> list[list[str]]:
    answer4, next_question = split_merged_segment(
        segments[4],
        r"(?=Guruh trener bellashuv vaqtida berilgan joyda Yoki gilam atrofidagi masofada turishi mumkin\?)",
    )
    return [
        [segments[0], segments[1], segments[2], segments[3], answer4],
        [next_question, segments[5], segments[6], segments[7], segments[8]],
    ]


def uz_override_469(segments: list[str]) -> list[list[str]]:
    compact = [segment for segment in segments if segment]
    return [compact]


def ru_override_3(segments: list[str]) -> list[list[str]]:
    question = f"{segments[0]} {segments[2]}"
    return [[question, segments[1], segments[3], segments[4], segments[5]]]


def ru_override_216(segments: list[str]) -> list[list[str]]:
    answer4, next_question = split_merged_segment(segments[4], r"(?=Что такое руки назад\?)")
    return [
        [segments[0], segments[1], segments[2], segments[3], answer4],
        [next_question, segments[5], segments[6], segments[7], segments[8]],
    ]


def ru_override_374(_: list[str]) -> list[list[str]]:
    return [[
        "Как завершается комплекс упражнений?",
        "#1-2-3 упражнение окончено",
        "1-23-3-34-5-6 упражнение окончено",
        "1-2-33-43-5-6-9-19 упражнение окончено",
        "1-23-13-44 остановить упражнение",
    ]]


def ru_override_380(segments: list[str]) -> list[list[str]]:
    answer4, next_question = split_merged_segment(segments[4], r"(?=Что означает слово гигиена\?)")
    return [
        [segments[0], segments[1], segments[2], segments[3], answer4],
        [next_question, segments[5], segments[6], segments[7], segments[8]],
    ]


def ru_override_386(segments: list[str]) -> list[list[str]]:
    return [split_first_segment_if_needed(segments)]


UZ_OVERRIDES: dict[int, OverrideHandler] = {
    100: uz_override_100,
    157: uz_override_157,
    361: uz_override_361,
    386: uz_override_386,
    463: uz_override_463,
    469: uz_override_469,
}


RU_OVERRIDES: dict[int, OverrideHandler] = {
    3: ru_override_3,
    216: ru_override_216,
    374: ru_override_374,
    380: ru_override_380,
    386: ru_override_386,
}


def parse_legacy(text: str, overrides: dict[int, OverrideHandler] | None = None) -> list[Question]:
    questions: list[Question] = []
    active_overrides = overrides or {}

    for block_index, block_lines in enumerate(split_blocks(text), start=1):
        segments = split_first_segment_if_needed(block_to_segments(block_lines))
        transformed_blocks = active_overrides.get(block_index, lambda current: [current])(segments)
        for transformed in transformed_blocks:
            questions.append(make_question(transformed))

    return questions


def parse_sequential(text: str) -> list[Question]:
    questions: list[Question] = []

    for block in split_raw_blocks(text):
        if len(block) == 1 and is_header_line(block[0]):
            continue

        tokens = [normalize_segment(line) for line in block if not is_separator_line(line, SEPARATOR_RE)]
        tokens = [token for token in tokens if not is_separator_line(token, BLOCK_SEPARATOR_RE)]
        tokens = [token for token in tokens if token and not is_header_line(token)]
        tokens = remove_duplicate_number_token(tokens)

        if not tokens:
            continue
        if len(tokens) % 5 != 0:
            raise ValueError(f"Sequential parser could not split block into 5-part questions: {tokens}")

        for chunk in chunk_segments(tokens):
            questions.append(make_question(chunk))

    return questions


PROFILE_HANDLERS: dict[str, Callable[[str], list[Question]]] = {
    "legacy": lambda text: parse_legacy(text, {}),
    "sequential": parse_sequential,
    "jismoniy_uz": lambda text: parse_legacy(text, UZ_OVERRIDES),
    "jismoniy_ru": lambda text: parse_legacy(text, RU_OVERRIDES),
}


PROFILE_LABELS = {
    "auto": "Auto",
    "legacy": "Legacy",
    "sequential": "Sequential",
    "jismoniy_uz": "Jismoniy UZ",
    "jismoniy_ru": "Jismoniy RU",
}


def verify_questions(questions: list[Question]) -> None:
    for index, question in enumerate(questions, start=1):
        if len(question.answers) != 4:
            raise ValueError(f"Question {index} does not have 4 answers")
        if sum(answer.startswith("#") for answer in question.answers) != 1:
            raise ValueError(f"Question {index} does not have exactly one correct answer")
        if "====" in question.prompt or "++++" in question.prompt:
            raise ValueError(f"Question {index} still contains separators: {question.prompt}")


def render_hemis_text(questions: list[Question]) -> str:
    lines: list[str] = []
    for index, question in enumerate(questions, start=1):
        lines.append(f"{index}. {question.prompt}")
        for answer in question.answers:
            lines.append(f"===== {answer}")
        lines.append("+++++")
    return "\n".join(lines)


def get_profile_order(source_name: str, requested_profile: str) -> list[str]:
    if requested_profile != "auto":
        return [requested_profile]

    name = source_name.lower()
    if "jismoniy" in name and "rus" in name:
        return ["jismoniy_ru", "jismoniy_uz", "sequential", "legacy"]
    if "jismoniy" in name:
        return ["jismoniy_uz", "jismoniy_ru", "sequential", "legacy"]
    if "xorijiy" in name or "tillar" in name or "ingliz" in name or "deutsch" in name:
        return ["sequential", "legacy", "jismoniy_uz", "jismoniy_ru"]
    return ["sequential", "legacy", "jismoniy_uz", "jismoniy_ru"]


def get_encoding_order(source_name: str, requested_encoding: str) -> list[str]:
    if requested_encoding != "auto":
        return [requested_encoding]

    name = source_name.lower()
    if "rus" in name or "russian" in name:
        return ["cp1251", "cp1254", "utf-8-sig", "utf-8", "cp1252", "latin1"]
    return list(ENCODING_CANDIDATES)


def choose_candidate(raw_bytes: bytes, source_name: str, profile: str, encoding: str) -> tuple[str, str, list[Question]]:
    errors: list[str] = []
    best_choice: tuple[int, int, int, str, str, list[Question]] | None = None

    profile_order = get_profile_order(source_name, profile)
    encoding_order = get_encoding_order(source_name, encoding)

    for profile_rank, profile_name in enumerate(profile_order):
        handler = PROFILE_HANDLERS[profile_name]
        for encoding_rank, encoding_name in enumerate(encoding_order):
            try:
                text = raw_bytes.decode(encoding_name, errors="replace")
                questions = handler(text)
                verify_questions(questions)
            except Exception as exc:
                errors.append(f"{profile_name}/{encoding_name}: {exc}")
                continue

            score = (len(questions), -profile_rank, -encoding_rank)
            if best_choice is None or score > best_choice[:3]:
                best_choice = (*score, profile_name, encoding_name, questions)

    if best_choice is None:
        sample = "\n".join(errors[:6])
        raise ValueError(f"Faylni formatlab bo'lmadi.\n{sample}")

    _, _, _, profile_name, encoding_name, questions = best_choice
    return profile_name, encoding_name, questions


def call_with_retry(action: Callable[[], object], attempts: int = 6, delay: float = 0.5):
    last_error: Exception | None = None
    retry_errors = [AttributeError]
    if pywintypes is not None:
        retry_errors.append(pywintypes.com_error)

    for _ in range(attempts):
        try:
            return action()
        except tuple(retry_errors) as exc:
            last_error = exc
            time.sleep(delay)
    if last_error is not None:
        raise last_error
    raise RuntimeError("Unknown Word automation failure")


@contextmanager
def word_application():
    if pythoncom is None or win32 is None:
        raise RuntimeError("Microsoft Word automation topilmadi.")
    pythoncom.CoInitialize()
    word = None
    try:
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        yield word
    finally:
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def export_word_to_text(word_app, source_path: Path, temp_path: Path) -> None:
    document = call_with_retry(
        lambda: word_app.Documents.Open(
            FileName=str(source_path.resolve()),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
        )
    )
    try:
        call_with_retry(
            lambda: document.SaveAs2(
                FileName=str(temp_path.resolve()),
                FileFormat=WORD_SAVE_AS_TEXT,
            )
        )
    finally:
        call_with_retry(lambda: document.Close(False))


def write_word_document(output_path: Path, content: str) -> None:
    document = Document()
    lines = content.splitlines() or [""]

    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    document.save(output_path)


def make_output_stem(source_name: str) -> str:
    stem = Path(source_name).stem.strip()
    return re.sub(r"_hemis$", "", stem, flags=re.IGNORECASE)


def find_sidecar_text_paths(source_path: Path) -> list[Path]:
    candidates = [
        source_path.with_suffix(".txt"),
        source_path.with_name(f"{source_path.stem}.export.txt"),
        source_path.with_name(f"{source_path.stem}_hemis.txt"),
    ]
    unique_paths: list[Path] = []
    for candidate in candidates:
        if candidate.exists() and candidate not in unique_paths:
            unique_paths.append(candidate)
    return unique_paths


def _required_office_backend_error() -> str:
    system = platform.system()
    if system == "Windows":
        return (
            "Windows serverda Word ham, LibreOffice ham topilmadi. "
            "Generatsiya uchun Microsoft Word yoki LibreOffice o'rnating. "
            "Masalan: winget install TheDocumentFoundation.LibreOffice"
        )
    if system == "Linux":
        return (
            "Linux serverda LibreOffice topilmadi. "
            "Generatsiya uchun LibreOffice o'rnating. "
            "Masalan: sudo apt update && sudo apt install -y libreoffice"
        )
    if system == "Darwin":
        return (
            "macOS serverda LibreOffice topilmadi. "
            "Generatsiya uchun LibreOffice o'rnating. "
            "Masalan: brew install --cask libreoffice"
        )
    return "Serverda generatsiya uchun kerakli office backend topilmadi."


def _find_soffice() -> Path | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if platform.system() == "Windows" and not soffice:
        for candidate in (
            Path("C:/Program Files/LibreOffice/program/soffice.exe"),
            Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),
        ):
            if candidate.exists():
                return candidate
    return Path(soffice) if soffice else None


def export_libreoffice_to_text(source_path: Path, temp_path: Path, soffice_path: Path) -> None:
    temp_path.parent.mkdir(parents=True, exist_ok=True)
    work_source = temp_path.parent / f"source{source_path.suffix.lower()}"
    if work_source.exists():
        work_source.unlink()
    shutil.copy2(source_path, work_source)
    source_for_text = work_source

    if work_source.suffix.lower() == ".doc":
        docx_proc = subprocess.run(
            [str(soffice_path), "--headless", "--convert-to", "docx", "--outdir", str(temp_path.parent), str(work_source.resolve())],
            capture_output=True,
            text=True,
            check=False,
        )
        converted_docx = temp_path.parent / "source.docx"
        if not converted_docx.exists():
            raise RuntimeError(docx_proc.stderr.strip() or docx_proc.stdout.strip() or "LibreOffice DOCX konvertatsiyasida xatolik berdi.")
        source_for_text = converted_docx

    proc = subprocess.run(
        [str(soffice_path), "--headless", "--convert-to", "txt:Text", "--outdir", str(temp_path.parent), str(source_for_text.resolve())],
        capture_output=True,
        text=True,
        check=False,
    )
    if proc.returncode != 0 and not any(temp_path.parent.glob("*.txt")):
        raise RuntimeError(proc.stderr.strip() or proc.stdout.strip() or "LibreOffice matn eksportida xatolik berdi.")

    converted_path = temp_path.parent / f"{source_path.stem}.txt"
    if not converted_path.exists():
        txt_files = sorted(temp_path.parent.glob("*.txt"))
        if not txt_files:
            raise RuntimeError("LibreOffice TXT fayl yaratmadi.")
        converted_path = txt_files[0]

    if converted_path.resolve() != temp_path.resolve():
        if temp_path.exists():
            temp_path.unlink()
        shutil.move(str(converted_path), str(temp_path))


def export_source_to_text(source_path: Path, temp_path: Path) -> str:
    errors: list[str] = []

    if platform.system() == "Windows":
        try:
            with word_application() as word:
                export_word_to_text(word, source_path, temp_path)
            return "word"
        except Exception as exc:
            errors.append(f"Word: {exc}")

    soffice_path = _find_soffice()
    if soffice_path is not None:
        try:
            export_libreoffice_to_text(source_path, temp_path, soffice_path)
            return "libreoffice"
        except Exception as exc:
            errors.append(f"LibreOffice: {exc}")

    detail = f" {'; '.join(errors[:2])}" if errors else ""
    raise RuntimeError(_required_office_backend_error() + detail)


def format_word_file(
    source_path: Path,
    output_dir: Path | None = None,
    profile: str = "auto",
    encoding: str = "auto",
) -> ConversionResult:
    source_path = Path(source_path)
    if not source_path.exists():
        raise FileNotFoundError(source_path)

    target_dir = Path(output_dir) if output_dir else source_path.parent
    target_dir.mkdir(parents=True, exist_ok=True)
    output_stem = make_output_stem(source_path.name)
    txt_path = target_dir / f"{output_stem}_hemis.txt"
    docx_path = target_dir / f"{output_stem}_hemis.docx"

    with tempfile.TemporaryDirectory() as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        temp_txt = temp_dir / f"{output_stem}.txt"
        raw_candidates: list[tuple[str, bytes]] = []

        try:
            backend = export_source_to_text(source_path, temp_txt)
            raw_candidates.append((backend, temp_txt.read_bytes()))
        except Exception as exc:
            backend_error = exc
        else:
            backend_error = None

        for sidecar_path in find_sidecar_text_paths(source_path):
            raw_candidates.append((f"sidecar:{sidecar_path.name}", sidecar_path.read_bytes()))

        if not raw_candidates:
            raise backend_error or RuntimeError("Fayl uchun xom matn manbasi topilmadi.")

        parse_errors: list[str] = []
        selected_profile = ""
        selected_encoding = ""
        questions: list[Question] | None = None
        backend = raw_candidates[0][0]

        for candidate_backend, raw_bytes in raw_candidates:
            try:
                selected_profile, selected_encoding, questions = choose_candidate(
                    raw_bytes,
                    source_path.name,
                    profile,
                    encoding,
                )
                backend = candidate_backend
                break
            except Exception as exc:
                parse_errors.append(f"{candidate_backend}: {exc}")

        if questions is None:
            raise ValueError("\n".join(parse_errors[:3]))

        hemis_text = render_hemis_text(questions)
        txt_path.write_text(hemis_text, encoding="utf-8")
        write_word_document(docx_path, hemis_text)

    return ConversionResult(
        source_name=source_path.name,
        profile=selected_profile,
        encoding=selected_encoding,
        backend=backend,
        question_count=len(questions),
        hemis_text=hemis_text,
        txt_path=txt_path,
        docx_path=docx_path,
    )


def format_uploaded_bytes(
    source_name: str,
    content: bytes,
    output_dir: Path,
    profile: str = "auto",
    encoding: str = "auto",
) -> ConversionResult:
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory() as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        source_path = temp_dir / source_name
        source_path.write_bytes(content)
        return format_word_file(source_path, output_dir=output_dir, profile=profile, encoding=encoding)


def preview_text(hemis_text: str, question_limit: int = 5) -> str:
    lines = hemis_text.splitlines()
    return "\n".join(lines[: question_limit * 6])
