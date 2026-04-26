import csv
import json
import random
import tempfile
import os
import time
import uuid
import gzip
from pathlib import Path
from io import BytesIO
import base64
import shutil
import subprocess
import mimetypes
import html as std_html
import zipfile
from urllib.parse import unquote
from django.contrib.auth.decorators import login_required
from django.core.files.base import ContentFile
from django.http import FileResponse, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404
from django.db import transaction
from django.db.models import Prefetch
from django.utils.decorators import method_decorator
from django.views import View
from docx.shared import Inches
from openpyxl import load_workbook, Workbook
import re
from docx import Document
from sympy import sympify, latex
import matplotlib
from lxml import html as lxml_html
from PIL import Image

matplotlib.use("Agg")
import matplotlib.pyplot as plt

DOCX_XMLNS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "o": "urn:schemas-microsoft-com:office:office",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v": "urn:schemas-microsoft-com:vml",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}

BROWSER_SAFE_IMAGE_MIME_TYPES = {
    "image/apng",
    "image/avif",
    "image/bmp",
    "image/gif",
    "image/jpeg",
    "image/png",
    "image/svg+xml",
    "image/webp",
}

IMPORT_SEPARATOR_RE = re.compile(r"^(?:\++|-{3,})$")
IMPORT_ANSWER_PREFIX_RE = re.compile(r"^=+\s*")
IMPORT_CORRECT_PREFIX_RE = re.compile(r"^#+\s*")
REPO_ROOT = Path(__file__).resolve().parents[3]
WORD_HELPER_BRIDGE_DIR = REPO_ROOT / "desktop" / "windows_word_bridge"
WORD_HELPER_EXTENSION_DIR = REPO_ROOT / "desktop" / "browser_extension" / "quiz_word_helper"
OFFICE_HELPER_EXE_PATH = WORD_HELPER_BRIDGE_DIR / "dist" / "QuizOfficeHelper.exe"


def _local_name(tag):
    return tag.split("}", 1)[1] if "}" in tag else tag


def _normalize_text(text):
    return " ".join((text or "").split())


def _docx_convert_image_to_png(blob):
    try:
        image = Image.open(BytesIO(blob))
        if image.mode not in ("RGB", "RGBA"):
            image = image.convert("RGBA")
        buffer = BytesIO()
        image.save(buffer, format="PNG")
        return buffer.getvalue()
    except Exception:
        return None


def _docx_media_map(doc):
    media = {}
    for rel_id, rel in doc.part.rels.items():
        if "image" not in rel.reltype:
            continue

        part = getattr(rel, "target_part", None)
        if part is None:
            continue

        target_ref = getattr(rel, "target_ref", "") or ""
        blob = part.blob
        lower_target = target_ref.lower()

        if lower_target.endswith((".emz", ".wmz")):
            try:
                blob = gzip.decompress(blob)
                target_ref = target_ref[:-4] + (".emf" if lower_target.endswith(".emz") else ".wmf")
            except OSError:
                pass

        guessed_mime = mimetypes.guess_type(target_ref)[0]
        mime = guessed_mime or getattr(part, "content_type", None) or "application/octet-stream"

        if mime not in BROWSER_SAFE_IMAGE_MIME_TYPES:
            converted = _docx_convert_image_to_png(blob)
            if converted is None:
                continue
            blob = converted
            mime = "image/png"

        b64 = base64.b64encode(blob).decode("ascii")
        media[rel_id] = f"data:{mime};base64,{b64}"
    return media


def _docx_formula_to_html(node):
    text = _normalize_text("".join(node.itertext()))
    if not text:
        text = "[formula]"
    return f'<span class="docx-formula">{std_html.escape(text)}</span>'


def _docx_collect_image_html(node, media_map):
    html_parts = []

    rel_embed = f"{{{DOCX_XMLNS['r']}}}embed"
    rel_id_attr = f"{{{DOCX_XMLNS['r']}}}id"
    rel_legacy_attr = f"{{{DOCX_XMLNS['o']}}}relid"

    for child in node.iter():
        tag = _local_name(child.tag)
        rel_id = None

        if tag == "blip":
            rel_id = child.get(rel_embed)
        elif tag == "imagedata":
            rel_id = child.get(rel_id_attr) or child.get(rel_legacy_attr)

        if not rel_id:
            continue

        src = media_map.get(rel_id)
        if src and "image/x-wmf" in src:
            import base64
            # We deferred this. Now convert it using proper dimensions!
            raw_blob = base64.b64decode(src.split(",")[1])
            converted = _docx_convert_image_to_png(raw_blob, width_px, height_px)
            if converted:
                encoded = base64.b64encode(converted).decode("ascii")
                src = f"data:image/png;base64,{encoded}"
                
        if src:
            html_parts.append(f'<img src="{src}" alt="image">')
        else:
            html_parts.append('<span class="docx-image-placeholder">[image]</span>')

    return html_parts


def _docx_render_inline(node, media_map):
    tag = _local_name(node.tag)

    if tag in {"t", "instrText", "delText"}:
        return std_html.escape(node.text or "")
    if tag in {"tab"}:
        return "    "
    if tag in {"br", "cr"}:
        return "<br>"
    if tag in {"noBreakHyphen"}:
        return "-"
    if tag == "sym":
        char = node.get(f"{{{DOCX_XMLNS['w']}}}char") or ""
        try:
            return chr(int(char, 16)) if char else ""
        except ValueError:
            return ""
    if tag in {"drawing", "pict", "object"}:
        return "".join(_docx_collect_image_html(node, media_map))
    if tag in {"oMath", "oMathPara"}:
        return _docx_formula_to_html(node)
    if tag in {"r", "smartTag", "sdtContent", "ins", "del"}:
        return "".join(_docx_render_inline(child, media_map) for child in node)
    if tag == "hyperlink":
        inner = "".join(_docx_render_inline(child, media_map) for child in node)
        if not inner:
            return ""
        return f"<span>{inner}</span>"

    return "".join(_docx_render_inline(child, media_map) for child in node)


def _docx_paragraph_to_html(paragraph, media_map):
    inner = "".join(_docx_render_inline(child, media_map) for child in paragraph._p)
    inner = inner.strip()

    if not inner:
        text = std_html.escape(paragraph.text or "")
        inner = text.strip()

    return f"<p>{inner}</p>" if inner else ""


def _docx_table_to_html(table, media_map):
    rows_html = []

    for row in table.rows:
        cells_html = []
        for cell in row.cells:
            paragraph_html = []
            for paragraph in cell.paragraphs:
                rendered = _docx_paragraph_to_html(paragraph, media_map)
                if rendered:
                    paragraph_html.append(rendered)

            cell_inner = "".join(paragraph_html).strip()
            if not cell_inner:
                cell_text = _normalize_text(cell.text)
                if cell_text:
                    cell_inner = f"<p>{std_html.escape(cell_text)}</p>"

            cells_html.append(f"<td>{cell_inner}</td>")

        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")

    return "<table>" + "".join(rows_html) + "</table>" if rows_html else ""


def _docx_to_html_embedded_pure(docx_path):
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(docx_path)
    media_map = _docx_media_map(doc)

    blocks = []
    for child in doc.element.body.iterchildren():
        tag = _local_name(child.tag)

        if tag == "p":
            rendered = _docx_paragraph_to_html(Paragraph(child, doc), media_map)
            if rendered:
                blocks.append(rendered)
        elif tag == "tbl":
            rendered = _docx_table_to_html(Table(child, doc), media_map)
            if rendered:
                blocks.append(rendered)

    return "<html><body>" + "".join(blocks) + "</body></html>"


IMPORT_TEXT_REPLACEMENTS = {
    "\uf025": "%",
    "\uf0b0": "°",
    "\uf0b1": "±",
    "\uf0d7": "×",
    "\uf0f7": "÷",
    "вЂ˜": "'",
    "вЂ™": "'",
    "вЂњ": '"',
    "вЂќ": '"',
    "вЂ“": "-",
    "вЂ”": "-",
    "вЂ¦": "...",
    "Â°": "°",
    "Â±": "±",
    "Â·": "·",
}


def _cleanup_import_html(value):
    cleaned = value or ""
    for old, new in IMPORT_TEXT_REPLACEMENTS.items():
        cleaned = cleaned.replace(old, new)
    return cleaned


IMPORT_ALLOWED_TAGS = {
    "a",
    "b",
    "blockquote",
    "br",
    "div",
    "em",
    "h1",
    "h2",
    "h3",
    "h4",
    "hr",
    "i",
    "img",
    "li",
    "ol",
    "p",
    "span",
    "strong",
    "sub",
    "sup",
    "table",
    "tbody",
    "td",
    "th",
    "thead",
    "tr",
    "u",
    "ul",
}

IMPORT_ALLOWED_ATTRIBUTES = {
    "a": {"href", "rel", "target"},
    "img": {"alt", "src"},
    "span": {"class"},
    "td": {"colspan", "rowspan"},
    "th": {"colspan", "rowspan"},
}


def _element_has_meaningful_content(element):
    text = _cleanup_import_html("".join(element.itertext()).replace("\xa0", " ")).strip()
    if text:
        return True
    return bool(element.xpath(".//img | .//table | .//br"))


def _normalize_import_fragment(value):
    cleaned = _cleanup_import_html(value).replace("\xa0", " ").strip()
    if not cleaned:
        return ""

    try:
        root = lxml_html.fromstring(f"<div>{cleaned}</div>")
    except Exception:
        return std_html.escape(cleaned)

    for node in root.xpath(".//comment()"):
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)

    for element in list(root.iter())[1:]:
        if not isinstance(element.tag, str):
            continue

        tag = _local_name(element.tag).lower()
        style = (element.get("style") or "").lower()

        if tag in {"script", "style", "meta", "link", "title", "head"}:
            element.drop_tree()
            continue

        if tag in {"html", "body"}:
            element.drop_tag()
            continue

        if tag in {"font", "span"}:
            classes = set(filter(None, (element.get("class") or "").split()))
            if "vertical-align: super" in style:
                element.tag = "sup"
                tag = "sup"
            elif "vertical-align: sub" in style:
                element.tag = "sub"
                tag = "sub"
            elif tag == "span" and "docx-formula" in classes:
                element.attrib.clear()
                element.set("class", "docx-formula")
                continue
            else:
                element.drop_tag()
                continue

        if tag not in IMPORT_ALLOWED_TAGS:
            element.drop_tag()
            continue

        allowed_attrs = IMPORT_ALLOWED_ATTRIBUTES.get(tag, set())
        for attr in list(element.attrib):
            if attr.lower() not in allowed_attrs:
                del element.attrib[attr]

        if tag == "img":
            src = (element.get("src") or "").strip()
            if not src or src.lower().startswith("javascript:"):
                element.drop_tree()
                continue
            element.set("alt", _cleanup_import_html(element.get("alt") or "image"))
        elif tag == "a":
            href = (element.get("href") or "").strip()
            if href.lower().startswith("javascript:"):
                element.drop_tag()
                continue
            if href:
                element.set("href", href)
                element.set("target", "_blank")
                element.set("rel", "noopener noreferrer")

    for element in list(root.iter())[::-1]:
        if element is root or not isinstance(element.tag, str):
            continue
        tag = _local_name(element.tag).lower()
        if tag in {"p", "div", "li", "td", "th"} and not _element_has_meaningful_content(element):
            element.drop_tree()

    parts = []
    if (root.text or "").strip():
        parts.append(std_html.escape(_cleanup_import_html(root.text)))
    for child in root:
        parts.append(lxml_html.tostring(child, encoding="unicode", method="html"))

    normalized = _cleanup_import_html("".join(parts)).strip()
    normalized = re.sub(r"(?:\s*\r?\n\s*){3,}", "\n\n", normalized)
    return normalized


def clean_import_question_fragment(value):
    return _normalize_import_fragment(value)


def clean_import_answer_fragment(value):
    return _normalize_import_fragment(value)


def clean_import_question_text(value):
    return clean_import_question_fragment(value)


def clean_import_answer_text(value):
    return clean_import_answer_fragment(value)


def _iter_import_blocks(el):
    for child in el.iterchildren():
        tag = str(child.tag).lower()
        if "}" in tag:
            tag = tag.split("}", 1)[1]

        if tag in ("p", "h1", "h2", "h3", "h4", "table", "div", "li"):
            yield child
        elif tag in ("ul", "ol"):
            for li in child.xpath("./*[local-name()='li']"):
                yield li
        else:
            yield from _iter_import_blocks(child)


def _coerce_import_questions(raw_questions):
    normalized_questions = []

    for question in raw_questions or []:
        if not isinstance(question, dict):
            continue

        question_text = question.get("text") or ""
        image_base64 = question.get("image_base64")
        answers = []

        for answer in question.get("answers") or []:
            if isinstance(answer, dict):
                answer_text = answer.get("text") or ""
                is_correct = bool(answer.get("is_correct"))
            elif isinstance(answer, (list, tuple)) and len(answer) >= 2:
                answer_text = answer[0] or ""
                is_correct = bool(answer[1])
            else:
                continue

            answers.append({
                "text": answer_text,
                "is_correct": is_correct,
            })

        normalized_questions.append({
            "text": question_text,
            "image_base64": image_base64,
            "answers": answers,
        })

    return normalized_questions


def _strip_import_prefix(fragment_html, *patterns):
    fragment_html = _normalize_import_fragment(fragment_html or "")
    if not fragment_html or not patterns:
        return fragment_html

    try:
        root = lxml_html.fromstring(f"<div>{fragment_html}</div>")
    except Exception:
        return fragment_html

    def strip_first_match(value):
        if not value:
            return value

        leading = value[: len(value) - len(value.lstrip())]
        stripped = value.lstrip()
        original = stripped

        for pattern in patterns:
            stripped = pattern.sub("", stripped, count=1)

        if stripped == original:
            return value
        return f"{leading}{stripped}"

    for node in root.iter():
        if node.text and node.text.strip():
            node.text = strip_first_match(node.text)
            break

    return _normalize_import_fragment(
        "".join(
            lxml_html.tostring(child, encoding="unicode", method="html")
            for child in root
        )
    )


def _is_import_separator(value):
    return bool(IMPORT_SEPARATOR_RE.match((value or "").strip()))


def _parse_prefixed_import_answer(compact_text, block_html):
    compact_text = (compact_text or "").strip()
    if not compact_text:
        return None

    prefix_match = IMPORT_ANSWER_PREFIX_RE.match(compact_text)
    if not prefix_match:
        return None

    answer_text = compact_text[prefix_match.end():].strip()
    if not answer_text:
        return None
    is_correct = answer_text.startswith("#")

    clean_html = _strip_import_prefix(block_html, IMPORT_ANSWER_PREFIX_RE)
    if is_correct:
        clean_html = _strip_import_prefix(clean_html, IMPORT_CORRECT_PREFIX_RE)

    return clean_html, is_correct


def _is_answer_prefix_only(compact_text):
    compact_text = (compact_text or "").strip()
    if not compact_text:
        return False

    prefix_match = IMPORT_ANSWER_PREFIX_RE.match(compact_text)
    if not prefix_match:
        return False

    return not compact_text[prefix_match.end():].strip()


def _extract_table_fallback_questions(body):
    tables = body.xpath(".//*[local-name()='table']")
    out2 = []
    for table in tables:
        for tr in table.xpath(".//tr"):
            cells = tr.xpath("./th|./td")
            if len(cells) < 3:
                continue
            
            stem_cell = cells[0]
            option_cells = cells[1:5]
            
            def cell_to_html(cell):
                parts = []
                for child in cell:
                    parts.append(_normalize_import_fragment(
                        lxml_html.tostring(child, encoding="unicode", method="html")
                    ))
                inner = "".join(parts).strip()
                if not inner:
                    txt = " ".join((cell.text_content() or "").split())
                    if txt:
                        inner = _normalize_import_fragment(f"<p>{std_html.escape(_cleanup_import_html(txt))}</p>")
                return inner if inner else ""
            
            q_html = cell_to_html(stem_cell)
            if not q_html: continue
            
            answers = []
            last_txt = None
            for c in option_cells:
                c_html = cell_to_html(c)
                if c_html and c_html != last_txt:
                    answers.append((c_html, False))
                    last_txt = c_html
                    
            if len(answers) < 2:
                continue
                
            answers[0] = (answers[0][0], True)
            out2.append({
                "text": q_html,
                "image_base64": None,
                "answers": answers
            })

    return out2


def _extract_import_questions_from_body(body):
    q_re = re.compile(r"^Savol\s*(\d+)\s*[:.)]?\s*", re.IGNORECASE)
    q_label_only_re = re.compile(r"^Savol\s*\d+\s*[:.)]?\s*$", re.IGNORECASE)
    v_re = re.compile(r"^Variant\s*(\d+)\b[\s:.)-]*", re.IGNORECASE)

    questions = []
    current_q_text = []
    current_answers = []
    waiting_for_answer_content = False

    def flush_current_question():
        nonlocal current_q_text, current_answers
        if current_q_text and current_answers:
            questions.append({
                "text": "\n".join(current_q_text),
                "image_base64": None,
                "answers": current_answers,
            })
        current_q_text = []
        current_answers = []

    for block in _iter_import_blocks(body):
        text = _cleanup_import_html(" ".join((block.text_content() or "").split()))
        has_img = bool(
            block.xpath(".//*[local-name()='img']")
            or block.xpath(".//*[local-name()='math']")
            or block.xpath(".//*[local-name()='imagedata']")
        )
        if not text and not has_img:
            continue

        block_html = _normalize_import_fragment(
            lxml_html.tostring(block, encoding="unicode", method="html")
        )
        if not block_html and not has_img:
            continue

        compact_text = text.strip()

        if _is_import_separator(compact_text):
            flush_current_question()
            waiting_for_answer_content = False
            continue

        if q_label_only_re.match(compact_text):
            flush_current_question()
            waiting_for_answer_content = False
            continue

        if q_re.match(compact_text):
            flush_current_question()
            clean_question_html = _strip_import_prefix(block_html, q_re)
            if clean_question_html:
                current_q_text = [clean_question_html]
            waiting_for_answer_content = False
            continue

        if current_q_text:
            vm = v_re.match(compact_text)
            if vm:
                vnum = int(vm.group(1))
                t_low = compact_text.lower()
                is_correct = (vnum == 1) or ("to'g'ri" in t_low) or ("togri" in t_low)
                clean_answer_html = _strip_import_prefix(block_html, v_re)
                current_answers.append((clean_answer_html, is_correct))
                waiting_for_answer_content = False
                continue

            parsed_answer = _parse_prefixed_import_answer(compact_text, block_html)
            if parsed_answer:
                current_answers.append(parsed_answer)
                waiting_for_answer_content = False
                continue

            if _is_answer_prefix_only(compact_text):
                waiting_for_answer_content = True
                continue

            if waiting_for_answer_content:
                is_correct = bool(IMPORT_CORRECT_PREFIX_RE.match(compact_text))
                clean_html = block_html
                if is_correct:
                    clean_html = _strip_import_prefix(clean_html, IMPORT_CORRECT_PREFIX_RE)
                if clean_html:
                    current_answers.append((clean_html, is_correct))
                waiting_for_answer_content = False
                continue

        current_q_text.append(block_html)
        waiting_for_answer_content = False

    flush_current_question()
    if questions:
        return questions

    return _extract_table_fallback_questions(body)


def _build_import_html_from_text(raw_text):
    paragraphs = []
    for line in (raw_text or "").splitlines():
        clean_line = line.strip()
        if clean_line:
            paragraphs.append(f"<p>{std_html.escape(clean_line)}</p>")
    return "".join(paragraphs)


def _extract_import_questions_from_html(raw_html):
    raw_html = (raw_html or "").strip()
    if not raw_html:
        return []

    try:
        root = lxml_html.fromstring(raw_html)
    except Exception:
        root = lxml_html.fromstring(f"<div>{raw_html}</div>")

    body_list = root.xpath("//body")
    body = body_list[0] if body_list else root
    return _extract_import_questions_from_body(body)


def _score_import_questions(questions):
    total_answers = 0
    total_question_length = 0

    for question in questions or []:
        answers = question.get("answers") or []
        total_answers += len(answers)
        total_question_length += len(_cleanup_import_html(question.get("text") or ""))

    return (
        len(questions or []),
        total_answers,
        total_question_length,
    )


def parse_pasted_questions(pasted_html, pasted_text=""):
    raw_html = (pasted_html or "").strip()
    raw_text = (pasted_text or "").strip()
    if not raw_html and not raw_text:
        return []

    html_questions = _extract_import_questions_from_html(raw_html) if raw_html else []
    text_questions = _extract_import_questions_from_html(_build_import_html_from_text(raw_text)) if raw_text else []

    if _score_import_questions(text_questions) > _score_import_questions(html_questions):
        return text_questions
    if html_questions:
        return html_questions
    return text_questions


def has_meaningful_import_content(pasted_html="", pasted_text=""):
    raw_text = (pasted_text or "").replace("\xa0", " ").strip()
    if raw_text:
        return True

    raw_html = (pasted_html or "").strip()
    if not raw_html:
        return False

    try:
        root = lxml_html.fromstring(raw_html)
    except Exception:
        try:
            root = lxml_html.fromstring(f"<div>{raw_html}</div>")
        except Exception:
            return bool(_cleanup_import_html(raw_html))

    text_content = _cleanup_import_html(" ".join(root.itertext())).strip()
    if text_content:
        return True

    media_xpaths = (
        ".//*[local-name()='img']",
        ".//*[local-name()='table']",
        ".//*[local-name()='math']",
        ".//*[local-name()='imagedata']",
    )
    return any(root.xpath(xpath) for xpath in media_xpaths)


def _required_office_backend_error():
    import platform

    system = platform.system()
    if system == "Windows":
        return (
            "Windows serverda Word ham, LibreOffice ham topilmadi. "
            "Aniq import uchun Microsoft Word yoki LibreOffice o'rnating. "
            "Masalan: winget install TheDocumentFoundation.LibreOffice"
        )
    if system == "Linux":
        return (
            "Linux serverda LibreOffice topilmadi. "
            "Aniq import uchun LibreOffice o'rnating. "
            "Masalan: sudo apt update && sudo apt install -y libreoffice"
        )
    if system == "Darwin":
        return (
            "macOS serverda LibreOffice topilmadi. "
            "Aniq import uchun LibreOffice o'rnating. "
            "Masalan: brew install --cask libreoffice"
        )
    return (
        "Serverda kerakli office backend topilmadi. "
        "Aniq import uchun Microsoft Word yoki LibreOffice o'rnating."
    )


def _resolve_office_backend():
    import platform

    system = platform.system()
    soffice = shutil.which("soffice") or shutil.which("libreoffice")

    if system == "Windows" and not soffice:
        windows_candidates = [
            Path("C:/Program Files/LibreOffice/program/soffice.exe"),
            Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),
        ]
        for candidate in windows_candidates:
            if candidate.exists():
                soffice = str(candidate)
                break

    if system == "Windows":
        try:
            import win32com.client  # noqa: F401
            return "word", None
        except Exception:
            if soffice:
                return "libreoffice", soffice
            raise RuntimeError(_required_office_backend_error())

    if soffice:
        return "libreoffice", soffice

    raise RuntimeError(_required_office_backend_error())


def _word_doc_to_docx_via_com(doc_path):
    try:
        import win32com.client
    except Exception as e:
        raise RuntimeError("MS Word (win32com) topilmadi!") from e

    from pathlib import Path
    doc_path = Path(doc_path).resolve()
    out_path = doc_path.with_suffix(".docx")

    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(doc_path))
        doc.SaveAs(str(out_path), FileFormat=16)
    finally:
        try:
            if doc: doc.Close(False)
        except: pass
        try:
            if word: word.Quit()
        except: pass
    return out_path


def _word_doc_to_docx_via_libreoffice(doc_path, soffice_path):
    import subprocess
    from pathlib import Path

    doc_path = Path(doc_path).resolve()
    out_dir = doc_path.parent
    subprocess.run(
        [soffice_path, "--headless", "--convert-to", "docx", "--outdir", str(out_dir), str(doc_path)],
        capture_output=True, text=True, check=False
    )
    out_path = doc_path.with_suffix(".docx")
    if not out_path.exists():
        raise RuntimeError("LibreOffice konvertatsiyada xatolik berdi")
    return out_path


def _word_doc_to_docx(doc_path):
    backend, soffice_path = _resolve_office_backend()
    if backend == "word":
        return _word_doc_to_docx_via_com(doc_path)
    return _word_doc_to_docx_via_libreoffice(doc_path, soffice_path)


def _detect_charset(raw):
    import re
    head = raw[:5000].decode("ascii", errors="ignore")
    m = re.search(r"charset=([\w-]+)", head, flags=re.IGNORECASE)
    if m: return m.group(1).strip().strip("\"'")
    return "utf-8"


def _embed_word_html_images(root, html_dir, images_dir):
    import base64
    import mimetypes
    from pathlib import Path
    from urllib.parse import unquote

    def resolve_src(src):
        src = (src or "").strip()
        if not src or src.lower().startswith("data:"): return None
        src = unquote(src)
        if src.lower().startswith("file:///"): return Path(src[8:])
        if src.lower().startswith("file:"): return Path(src[5:])
        src_norm = src.replace("\\", "/")
        try:
            for c in [html_dir / src_norm, images_dir / src_norm, images_dir / Path(src_norm).name, html_dir / Path(src_norm).name]:
                if c.exists(): return c
        except Exception:
            pass
        return None

    def embed_attr(el, attr):
        src = el.get(attr) or ""
        fp = resolve_src(src)
        if fp and fp.exists():
            mime = mimetypes.guess_type(str(fp))[0] or "application/octet-stream"
            b64 = base64.b64encode(fp.read_bytes()).decode("ascii")
            el.set(attr, f"data:{mime};base64,{b64}")

    for img in root.xpath("//*[local-name()='img']"):
        embed_attr(img, "src")
    for vml in root.xpath("//*[local-name()='imagedata']"):
        if vml.get("src"): embed_attr(vml, "src")


def _docx_to_html_embedded_via_word(docx_path):
    try:
        import win32com.client
    except Exception as e:
        raise RuntimeError("win32com yo'q") from e

    import tempfile
    from pathlib import Path
    from lxml import html as lxml_html
    docx_path = Path(docx_path).resolve()

    with tempfile.TemporaryDirectory(prefix="quiz_word_html_") as td:
        td_path = Path(td)
        out_html = td_path / "input.html"

        word = None
        doc = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(str(docx_path))
            doc.SaveAs(str(out_html), FileFormat=10)
        finally:
            try:
                if doc: doc.Close(False)
            except: pass
            try:
                if word: word.Quit()
            except: pass

        if not out_html.exists():
            raise RuntimeError("Word HTML yaratmadi.")

        raw = out_html.read_bytes()
        html_text = raw.decode(_detect_charset(raw), errors="replace")
        root = lxml_html.fromstring(html_text)
        images_dir = out_html.parent / f"{out_html.stem}.files"
        _embed_word_html_images(root, out_html.parent, images_dir)
        return lxml_html.tostring(root, encoding="unicode", method="html")


def _docx_to_html_embedded_via_libreoffice(docx_path, soffice_path):
    import subprocess
    import tempfile
    from pathlib import Path
    from lxml import html as lxml_html

    docx_path = Path(docx_path).resolve()
    with tempfile.TemporaryDirectory(prefix="quiz_word_html_") as td:
        subprocess.run([soffice_path, "--headless", "--convert-to", "html", "--outdir", td, str(docx_path)], capture_output=True)
        out_html = Path(td) / f"{docx_path.stem}.html"
        if not out_html.exists():
            out_html = Path(td) / f"{docx_path.stem}.htm"
        if not out_html.exists():
            raise RuntimeError("LibreOffice HTML formatda muvaffaqiyatli saqlay olmadi.")

        raw = out_html.read_bytes()
        html_text = raw.decode(_detect_charset(raw), errors="replace")
        root = lxml_html.fromstring(html_text)
        _embed_word_html_images(root, out_html.parent, out_html.parent)
        return lxml_html.tostring(root, encoding="unicode", method="html")


def _docx_to_html_embedded(docx_path):
    try:
        backend, soffice_path = _resolve_office_backend()
    except Exception:
        return _docx_to_html_embedded_pure(docx_path)

    try:
        if backend == "word":
            return _docx_to_html_embedded_via_word(docx_path)
        return _docx_to_html_embedded_via_libreoffice(docx_path, soffice_path)
    except Exception:
        return _docx_to_html_embedded_pure(docx_path)


def parse_word(file) -> list[dict]:
    import tempfile
    from pathlib import Path
    from lxml import html as lxml_html

    name = getattr(file, "name", "") or ""
    suffix = Path(name).suffix.lower()
    if suffix not in [".docx", ".doc"]:
        raise ValueError("Faqat .docx yoki .doc fayllar qabul qilinadi")

    with tempfile.TemporaryDirectory(prefix="quiz_word_") as td:
        img_path = Path(td)
        in_path = img_path / f"input{suffix}"
        
        with open(in_path, "wb") as f:
            if hasattr(file, "chunks"):
                for chunk in file.chunks(): f.write(chunk)
            else:
                f.write(file.read())

        docx_path = in_path
        if suffix == ".doc":
            try:
                docx_path = _word_doc_to_docx(in_path)
            except Exception as e:
                raise RuntimeError(
                    ".doc fayl legacy format. Uni import qilish uchun Word yoki LibreOffice kerak. "
                    "Cross-platform import uchun .docx dan foydalaning."
                ) from e

        html_text = _docx_to_html_embedded(docx_path)
        root = lxml_html.fromstring(html_text)
        body_list = root.xpath("//body")
        if not body_list:
            return []
        body = body_list[0]
        return _extract_import_questions_from_body(body)

from apps.question.models import Test, Question, Answer


def parse_excel(file, additional_files=None):
    wb = load_workbook(file)
    sheet = wb.active
    questions = []

    def convert_formula_to_base64(formula):
        """LaTeX yoki matematik formulani rasmga aylantirib, Base64 ga kodlaydi."""
        try:
            # Formulani LaTeX formatiga oвЂtkazish
            expr = sympify(formula)
            latex_expr = latex(expr)

            # Rasm yaratish
            plt.figure(figsize=(4, 1))
            plt.text(0.5, 0.5, f"${latex_expr}$", fontsize=12, ha='center', va='center')
            plt.axis('off')

            # Rasmni Base64 ga aylantirish
            buffer = BytesIO()
            plt.savefig(buffer, format='png', bbox_inches='tight', dpi=100)
            plt.close()
            buffer.seek(0)
            image_base64 = base64.b64encode(buffer.read()).decode('utf-8')
            return image_base64
        except Exception as e:
            print(f"Formula konvertatsiyasida xato: {str(e)}")
            return None

    for row in sheet.iter_rows(min_row=2, values_only=True):
        # Qator boвЂsh yoki yetarli ustunlarga ega emasligini tekshirish
        if not row or len(row) < 2:  # Kamida savol + 1 javob boвЂlishi kerak
            print(f"Qator oвЂtkazib yuborildi: yetarli ustunlar yoвЂq (ustunlar soni: {len(row) if row else 0})")
            continue

        question_text = row[0]
        if not question_text:
            print("Qator oвЂtkazib yuborildi: savol matni boвЂsh")
            continue

        # Formula aniqlash (masalan, LaTeX yoki oddiy matematik ifoda)
        formula_match = re.search(r'\[Formula:\s*([^\]]+)\]', question_text)
        image_base64 = None
        if formula_match:
            formula = formula_match.group(1).strip()
            image_base64 = convert_formula_to_base64(formula)
            # Formula belgisini matndan olib tashlash
            question_text = re.sub(r'\[Formula:\s*[^\]]+\]', '', question_text).strip()

        # [Rasm: nom.jpg] belgisini olib tashlash (agar mavjud boвЂlsa)
        question_text = re.sub(r'\[Rasm:\s*[^\]]+\]', '', question_text).strip()

        # Javoblarni olish (dinamik son)
        answers = [(_cleanup_import_html(str(row[i])), i == 1) for i in range(1, len(row)) if row[i]]  # 1-variant toвЂgвЂri
        if len(answers) < 2:  # Kamida 1 toвЂgвЂri va 1 notoвЂgвЂri javob
            print(f"Qator oвЂtkazib yuborildi: yetarli javoblar yoвЂq (javoblar soni: {len(answers)})")
            continue

        # Javoblarni tasodifiy tartibda almashtirish
        random.shuffle(answers)

        questions.append({
            'text': _cleanup_import_html(str(question_text)),
            'image_base64': image_base64,
            'answers': answers
        })

    print(f"Jami {len(questions)} ta savol topildi")
    return questions


QUESTION_IMPORT_BATCH_SIZE = 250
QUESTION_IMPORT_PREVIEW_TTL_SECONDS = 60 * 60


def _get_question_import_preview_dir():
    directory = Path(tempfile.gettempdir()) / "quiz_django_question_imports"
    directory.mkdir(parents=True, exist_ok=True)
    return directory


def _cleanup_expired_question_import_previews():
    now = time.time()
    for preview_path in _get_question_import_preview_dir().glob("*.json"):
        try:
            if now - preview_path.stat().st_mtime > QUESTION_IMPORT_PREVIEW_TTL_SECONDS:
                preview_path.unlink(missing_ok=True)
        except OSError:
            continue


def _get_question_import_preview_path(token):
    safe_token = re.sub(r"[^a-zA-Z0-9_-]", "", token or "")
    return _get_question_import_preview_dir() / f"{safe_token}.json"


def _save_question_import_preview(payload):
    _cleanup_expired_question_import_previews()
    token = uuid.uuid4().hex
    preview_path = _get_question_import_preview_path(token)
    with preview_path.open("w", encoding="utf-8") as preview_file:
        json.dump(payload, preview_file, ensure_ascii=False)
    return token


def _load_question_import_preview(token):
    preview_path = _get_question_import_preview_path(token)
    if not preview_path.exists():
        raise ValueError("Import preview topilmadi yoki muddati tugagan.")

    with preview_path.open("r", encoding="utf-8") as preview_file:
        return json.load(preview_file)


def _delete_question_import_preview(token):
    try:
        _get_question_import_preview_path(token).unlink(missing_ok=True)
    except OSError:
        pass


def _build_question_import_signature(question_text, answers):
    normalized_question = _normalize_import_fragment(question_text or "").strip().lower()
    if not normalized_question:
        return None

    correct_answers = sorted(
        {
            _normalize_import_fragment(answer.get("text") or "").strip().lower()
            for answer in (answers or [])
            if answer.get("is_correct")
            and _normalize_import_fragment(answer.get("text") or "").strip()
        }
    )
    if not correct_answers:
        return None
    return normalized_question, tuple(correct_answers)


def _prepare_question_import_entries(questions):
    prepared_entries = []
    seen_signatures = set()

    for question in questions or []:
        question_text = _normalize_import_fragment(question.get("text") or "").strip()
        normalized_answers = []
        for answer in question.get("answers") or []:
            cleaned_answer = _normalize_import_fragment(answer.get("text") or "").strip()
            if cleaned_answer:
                normalized_answers.append(
                    {
                        "text": cleaned_answer,
                        "is_correct": bool(answer.get("is_correct")),
                    }
                )

        if not question_text or len(normalized_answers) < 2:
            continue

        signature = _build_question_import_signature(question_text, normalized_answers)
        if signature is None or signature in seen_signatures:
            continue

        seen_signatures.add(signature)
        prepared_entries.append(
            {
                "text": question_text,
                "image_base64": question.get("image_base64") or None,
                "answers": normalized_answers,
                "signature": signature,
            }
        )

    return prepared_entries


def _apply_correct_answer_indexes(questions, correct_answer_indexes):
    if not isinstance(correct_answer_indexes, list):
        return questions

    for question_index, answer_indexes in enumerate(correct_answer_indexes):
        if question_index >= len(questions):
            break
        question = questions[question_index]
        answers = question.get("answers") or []
        valid_indexes = {
            int(index)
            for index in (answer_indexes or [])
            if str(index).strip().isdigit()
        }
        for answer_index, answer in enumerate(answers):
            answer["is_correct"] = answer_index in valid_indexes

    return questions


def _bulk_save_import_questions(test, questions):
    prepared_entries = _prepare_question_import_entries(questions)
    if not prepared_entries:
        return {"saved_count": 0, "skipped_duplicates": 0}

    existing_signatures = set()
    existing_questions = (
        Question.objects.filter(test=test)
        .only("id", "text")
        .prefetch_related(
            Prefetch(
                "answers",
                queryset=Answer.objects.filter(is_correct=True).only("question_id", "text"),
            )
        )
    )
    for existing_question in existing_questions:
        signature = _build_question_import_signature(
            existing_question.text,
            [
                {"text": answer.text, "is_correct": True}
                for answer in existing_question.answers.all()
            ],
        )
        if signature:
            existing_signatures.add(signature)

    entries_to_create = [
        entry
        for entry in prepared_entries
        if entry["signature"] not in existing_signatures
    ]
    skipped_duplicates = len(prepared_entries) - len(entries_to_create)

    if not entries_to_create:
        return {"saved_count": 0, "skipped_duplicates": skipped_duplicates}

    with transaction.atomic():
        question_objects = [
            Question(test=test, text=entry["text"])
            for entry in entries_to_create
        ]
        created_questions = Question.objects.bulk_create(
            question_objects,
            batch_size=QUESTION_IMPORT_BATCH_SIZE,
        )

        if created_questions and any(question.pk is None for question in created_questions):
            recent_questions = list(
                Question.objects.filter(test=test)
                .order_by("-id")[: len(created_questions)]
            )
            recent_questions.reverse()
            for created_question, recent_question in zip(created_questions, recent_questions):
                created_question.pk = recent_question.pk

        answer_objects = []
        for entry, created_question in zip(entries_to_create, created_questions):
            if entry["image_base64"]:
                image_data = base64.b64decode(entry["image_base64"])
                created_question.image.save(
                    f"question_{created_question.pk}.png",
                    ContentFile(image_data),
                )

            for answer in entry["answers"]:
                answer_objects.append(
                    Answer(
                        question_id=created_question.pk,
                        text=answer["text"],
                        is_correct=answer["is_correct"],
                    )
                )

        if answer_objects:
            Answer.objects.bulk_create(
                answer_objects,
                batch_size=QUESTION_IMPORT_BATCH_SIZE * 4,
            )

    return {
        "saved_count": len(entries_to_create),
        "skipped_duplicates": skipped_duplicates,
    }


def _iter_helper_archive_files(base_dir):
    for file_path in base_dir.rglob("*"):
        if file_path.is_dir():
            continue
        if any(part in {"__pycache__", ".venv"} for part in file_path.parts):
            continue
        if file_path.suffix.lower() == ".pyc":
            continue
        yield file_path


def _zip_directory_response(archive_name, directory_mappings, extra_files=None):
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        for archive_root, base_dir in directory_mappings:
            if not base_dir.exists():
                continue
            for file_path in _iter_helper_archive_files(base_dir):
                relative_path = file_path.relative_to(base_dir)
                archive_path = Path(archive_root) / relative_path
                archive.write(file_path, str(archive_path).replace("\\", "/"))

        for archive_path, content in extra_files or []:
            archive.writestr(archive_path, content)

    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type="application/zip")
    response["Content-Disposition"] = f'attachment; filename="{archive_name}"'
    return response


def _word_helper_kit_readme():
    return (
        "Quiz Word yordamchisi paketi\n"
        "=============================\n\n"
        "1. ZIP faylni oching.\n"
        "2. windows_word_bridge\\start_bridge.bat faylini ikki marta bosing.\n"
        "3. Word ichida savollar yozilgan faylni ochiq qoldiring.\n"
        "4. Brauzerda test sahifasiga qayting va \"Hujjatlar ro'yxatini yangilash\" tugmasini bosing.\n\n"
        "Agar brauzer Word yordamchisi bilan to'g'ridan-to'g'ri ishlamasa:\n"
        "5. browser_extension\\quiz_word_helper papkasini Chrome yoki Edge brauzeriga qo'shing.\n"
        "6. Sahifani yangilang va qayta urinib ko'ring.\n"
    )


def _get_word_application():
    try:
        import pythoncom
        import win32com.client
    except Exception as exc:
        raise RuntimeError("MS Word bilan ishlash uchun win32com mavjud emas.") from exc

    pythoncom.CoInitialize()
    word = win32com.client.Dispatch("Word.Application")
    return word


def _build_open_word_document_key(index, document):
    document_name = str(getattr(document, "Name", "") or "")
    document_full_name = str(getattr(document, "FullName", "") or "")
    return f"{index}:{document_full_name or document_name}"


def _list_open_word_documents():
    word = _get_word_application()
    documents = []
    active_full_name = ""
    try:
        active_document = getattr(word, "ActiveDocument", None)
        if active_document is not None:
            active_full_name = str(getattr(active_document, "FullName", "") or "")
    except Exception:
        active_full_name = ""

    for index in range(1, int(word.Documents.Count) + 1):
        document = word.Documents(index)
        name = str(getattr(document, "Name", "") or "").strip()
        full_name = str(getattr(document, "FullName", "") or "").strip()
        documents.append(
            {
                "key": _build_open_word_document_key(index, document),
                "name": name or f"Word hujjati {index}",
                "path": full_name,
                "is_active": bool(full_name and full_name == active_full_name),
                "is_saved": bool(getattr(document, "Saved", True)),
            }
        )

    return documents


def _find_open_word_document(document_key):
    word = _get_word_application()
    for index in range(1, int(word.Documents.Count) + 1):
        document = word.Documents(index)
        if _build_open_word_document_key(index, document) == document_key:
            return document
    raise ValueError("Tanlangan Word hujjati topilmadi. Uni qayta tanlang.")


def _extract_questions_from_open_word_document(document_key):
    document = _find_open_word_document(document_key)
    temp_directory = tempfile.TemporaryDirectory(prefix="quiz_open_word_import_")
    temp_dir_path = Path(temp_directory.name)
    full_name = str(getattr(document, "FullName", "") or "")
    suffix = Path(full_name).suffix.lower() if full_name else ".docx"
    if suffix not in {".doc", ".docx"}:
        suffix = ".docx"

    temp_file_path = temp_dir_path / f"open_word_copy{suffix}"
    temp_clone = None

    try:
        try:
            document.SaveCopyAs(str(temp_file_path))
        except Exception:
            temp_clone = document.Application.Documents.Add()
            temp_clone.Range().FormattedText = document.Range().FormattedText
            temp_file_path = temp_file_path.with_suffix(".docx")
            temp_clone.SaveAs2(str(temp_file_path), FileFormat=16)

        with temp_file_path.open("rb") as temp_file:
            return parse_word(temp_file)
    finally:
        if temp_clone is not None:
            try:
                temp_clone.Close(False)
            except Exception:
                pass
        temp_directory.cleanup()


@method_decorator(login_required, name='dispatch')
class ImportQuestionsView(View):
    def post(self, request, test_id):
        test = get_object_or_404(Test, id=test_id)

        try:
            started_at = time.perf_counter()
            pasted_html = (request.POST.get("pasted_html") or "").strip()
            pasted_text = (request.POST.get("pasted_text") or "").strip()
            has_pasted_content = has_meaningful_import_content(pasted_html, pasted_text)

            if has_pasted_content:
                questions = parse_pasted_questions(pasted_html, pasted_text)
            else:
                return JsonResponse({
                    "success": False,
                    "errors": "Savollar matnini paste qiling",
                })

            questions = _coerce_import_questions(questions)
            if not questions:
                return JsonResponse({
                    "success": False,
                    "errors": "Savollar topilmadi. Formatni tekshirib qayta paste qiling.",
                })

            preview_token = _save_question_import_preview(
                {
                    "requested_by": request.user.id,
                    "test_id": test_id,
                    "source": "paste",
                    "questions": questions,
                }
            )
            parse_duration_ms = round((time.perf_counter() - started_at) * 1000, 1)

            return JsonResponse({
                "success": True,
                "questions": questions,
                "preview_token": preview_token,
                "total_questions": len(questions),
                "parse_duration_ms": parse_duration_ms,
                "source_label": "Paste orqali import",
                "message": "Savollar muvaffaqiyatli o'qildi va preview tayyorlandi!",
            })
        except Exception as e:
            return JsonResponse({"success": False, "errors": f"Importda xato: {str(e)}"})


@method_decorator(login_required, name='dispatch')
class OpenWordDocumentsView(View):
    def get(self, request, test_id):
        try:
            documents = _list_open_word_documents()
            return JsonResponse(
                {
                    "success": True,
                    "supported": True,
                    "documents": documents,
                }
            )
        except Exception as exc:
            return JsonResponse(
                {
                    "success": True,
                    "supported": False,
                    "documents": [],
                    "message": (
                        "Server foydalanuvchining kompyuteridagi ochiq Word oynalarini ko'ra olmaydi. "
                        "Ubuntu serverda tavsiya etilgan usul: Word'dan nusxa olib, Clipboarddan yuklash."
                    ),
                    "errors": f"Ochiq Word hujjatlarini o'qishda xatolik: {str(exc)}",
                },
            )


@method_decorator(login_required, name='dispatch')
class ImportOpenWordDocumentView(View):
    def post(self, request, test_id):
        test = get_object_or_404(Test, id=test_id)
        document_key = (request.POST.get("document_key") or "").strip()
        if not document_key:
            return JsonResponse(
                {
                    "success": False,
                    "errors": "Avval ochiq Word hujjatini tanlang.",
                },
                status=400,
            )

        try:
            started_at = time.perf_counter()
            questions = _coerce_import_questions(_extract_questions_from_open_word_document(document_key))
            if not questions:
                return JsonResponse(
                    {
                        "success": False,
                        "errors": "Tanlangan Word hujjatidan savollar topilmadi.",
                    }
                )

            preview_token = _save_question_import_preview(
                {
                    "requested_by": request.user.id,
                    "test_id": test_id,
                    "source": "open-word",
                    "document_key": document_key,
                    "questions": questions,
                }
            )
            parse_duration_ms = round((time.perf_counter() - started_at) * 1000, 1)

            return JsonResponse(
                {
                    "success": True,
                    "questions": questions,
                    "preview_token": preview_token,
                    "total_questions": len(questions),
                    "parse_duration_ms": parse_duration_ms,
                    "source_label": "Ochiq Word hujjatidan import",
                    "message": "Word hujjatidagi savollar muvaffaqiyatli o'qildi!",
                }
            )
        except Exception as exc:
            return JsonResponse(
                {
                    "success": False,
                    "errors": f"Word hujjatini import qilishda xatolik: {str(exc)}",
                },
                status=500,
            )


@method_decorator(login_required, name='dispatch')
class DownloadWordHelperKitView(View):
    def get(self, request):
        return _zip_directory_response(
            "quiz-word-helper-kit.zip",
            [
                ("windows_word_bridge", WORD_HELPER_BRIDGE_DIR),
                ("browser_extension/quiz_word_helper", WORD_HELPER_EXTENSION_DIR),
            ],
            extra_files=[("README.txt", _word_helper_kit_readme())],
        )


@method_decorator(login_required, name='dispatch')
class DownloadWordHelperExtensionView(View):
    def get(self, request):
        return _zip_directory_response(
            "quiz-word-browser-extension.zip",
            [("quiz_word_helper", WORD_HELPER_EXTENSION_DIR)],
        )


@method_decorator(login_required, name='dispatch')
class DownloadOfficeHelperExeView(View):
    def get(self, request):
        if not OFFICE_HELPER_EXE_PATH.exists():
            return JsonResponse(
                {
                    "success": False,
                    "message": "Yordamchi dastur hali tayyor emas. Administrator exe faylini build qilishi kerak.",
                },
                status=503,
            )

        return FileResponse(
            OFFICE_HELPER_EXE_PATH.open("rb"),
            as_attachment=True,
            filename=OFFICE_HELPER_EXE_PATH.name,
            content_type="application/octet-stream",
        )


@method_decorator(login_required, name='dispatch')
class DownloadTemplateView(View):
    def get(self, request):
        document = Document()
        for line in [
            "++++",
            "Savol matni shu yerda",
            "====",
            "#To'g'ri javob",
            "====",
            "Javob 1",
            "====",
            "Javob 2",
            "====",
            "Javob 3",
            "++++",
        ]:
            document.add_paragraph(line)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            content=buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=questions_template.docx'
        return response


@method_decorator(login_required, name='dispatch')
class SaveImportedQuestionsView(View):
    def post(self, request):
        payload = {}
        if request.body:
            try:
                payload = json.loads(request.body.decode("utf-8"))
            except (ValueError, UnicodeDecodeError):
                payload = {}

        preview_token = (payload.get("preview_token") or "").strip()
        correct_answer_indexes = payload.get("correct_answer_indexes") or []
        test_id = None
        questions = None

        if preview_token:
            try:
                preview_payload = _load_question_import_preview(preview_token)
            except ValueError as exc:
                return JsonResponse({"success": False, "errors": str(exc)}, status=400)

            if preview_payload.get("requested_by") != request.user.id:
                return JsonResponse({"success": False, "errors": "Bu import preview sizga tegishli emas."}, status=403)

            test_id = preview_payload.get("test_id")
            questions = _coerce_import_questions(preview_payload.get("questions"))
            questions = _apply_correct_answer_indexes(questions, correct_answer_indexes)
        else:
            test_id = request.session.get("test_id")
            incoming_questions = _coerce_import_questions(payload.get("questions"))
            questions = incoming_questions or _coerce_import_questions(
                request.session.get("imported_questions", [])
            )

        if not test_id:
            return JsonResponse({"success": False, "errors": "Test topilmadi"})

        if not questions:
            return JsonResponse({"success": False, "errors": "Saqlash uchun savollar topilmadi"})

        test = get_object_or_404(Test, id=test_id)

        try:
            request.session.pop("imported_questions", None)
            request.session.pop("test_id", None)
            request.session.modified = True
            save_summary = _bulk_save_import_questions(test, questions)

            if preview_token:
                _delete_question_import_preview(preview_token)

            skipped_duplicates = save_summary["skipped_duplicates"]
            extra_message = f" {skipped_duplicates} ta takror savol o'tkazib yuborildi." if skipped_duplicates else ""
            return JsonResponse({
                "success": True,
                "saved_count": save_summary["saved_count"],
                "skipped_duplicates": skipped_duplicates,
                "message": f"{save_summary['saved_count']} ta savol muvaffaqiyatli saqlandi!{extra_message}",
            })
        except Exception as e:
            print(f"Saqlashda xato: {str(e)}")
            return JsonResponse({"success": False, "errors": f"Saqlashda xato: {str(e)}"})


@method_decorator(login_required, name='dispatch')
class DownloadTemplateView(View):
    def get(self, request):
        document = Document()
        for line in [
            "++++",
            "Savol matni shu yerda",
            "====",
            "#To'g'ri javob",
            "====",
            "Javob 1",
            "====",
            "Javob 2",
            "====",
            "Javob 3",
            "++++",
        ]:
            document.add_paragraph(line)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            content=buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=questions_template.docx'
        return response


@method_decorator(login_required, name='dispatch')
class ExportQuestionsView(View):
    def get(self, request, test_id, format_type):
        test = get_object_or_404(Test, id=test_id)
        questions = Question.objects.filter(test=test).prefetch_related('answers')

        if format_type == 'json':
            return self.export_json(questions, test)
        elif format_type == 'xlsx':
            return self.export_excel(questions, test)
        elif format_type == 'csv':
            return self.export_csv(questions, test)
        elif format_type == 'docx':
            return self.export_docx(questions, test)
        else:
            return HttpResponse(status=400, content="NotoвЂgвЂri format turi")

    def export_json(self, questions, test):
        data = {
            'test_name': test.name,
            'questions': [
                {
                    'text': q.text,
                    'image': q.image.url if q.image else None,
                    'answers': [
                        {'text': a.text, 'is_correct': a.is_correct}
                        for a in q.answers.all()
                    ]
                }
                for q in questions
            ]
        }
        response = HttpResponse(
            content_type='application/json',
            content=json.dumps(data, ensure_ascii=False, indent=2)
        )
        response['Content-Disposition'] = f'attachment; filename="{test.name}_questions.json"'
        return response

    def export_excel(self, questions, test):
        wb = Workbook()
        ws = wb.active
        ws.title = "Questions"

        headers = ["Savol matni", "Variant 1", "Variant 2", "Variant 3", "Variant 4", "ToвЂgвЂri javob", "Rasm"]
        ws.append(headers)

        for q in questions:
            answers = list(q.answers.all())
            correct_answer = next((a.text for a in answers if a.is_correct), '')
            incorrect_answers = [a.text for a in answers if not a.is_correct][:3]
            row = [
                q.text,
                correct_answer,  # Variant 1 doim toвЂgвЂri
                incorrect_answers[0] if len(incorrect_answers) > 0 else '',
                incorrect_answers[1] if len(incorrect_answers) > 1 else '',
                incorrect_answers[2] if len(incorrect_answers) > 2 else '',
                "1 variant",
                q.image.name if q.image else ''
            ]
            ws.append(row)

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            content=buffer
        )
        response['Content-Disposition'] = f'attachment; filename="{test.name}_questions.xlsx"'
        return response

    def export_csv(self, questions, test):
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{test.name}_questions.csv"'

        writer = csv.writer(response, lineterminator='\n')
        writer.writerow(["Savol matni", "Variant 1", "Variant 2", "Variant 3", "Variant 4", "ToвЂgвЂri javob", "Rasm"])

        for q in questions:
            answers = list(q.answers.all())
            correct_answer = next((a.text for a in answers if a.is_correct), '')
            incorrect_answers = [a.text for a in answers if not a.is_correct][:3]
            row = [
                q.text,
                correct_answer,
                incorrect_answers[0] if len(incorrect_answers) > 0 else '',
                incorrect_answers[1] if len(incorrect_answers) > 1 else '',
                incorrect_answers[2] if len(incorrect_answers) > 2 else '',
                "1 variant",
                q.image.name if q.image else ''
            ]
            writer.writerow(row)

        return response

    def export_docx(self, questions, test):
        doc = Document()
        doc.add_heading(f"{test.name} - Savollar", 0)

        for idx, q in enumerate(questions, 1):
            doc.add_paragraph(f"Savol {idx}: {q.text}")
            if q.image:
                try:
                    with q.image.open('rb') as img_file:
                        image_data = base64.b64encode(img_file.read()).decode('utf-8')
                    buffer = BytesIO(base64.b64decode(image_data))
                    doc.add_picture(buffer, width=Inches(4))
                except Exception:
                    doc.add_paragraph(f"[Rasm: {q.image.name}]")

            answers = list(q.answers.all())
            correct_answer = next((a.text for a in answers if a.is_correct), '')
            incorrect_answers = [a.text for a in answers if not a.is_correct][:3]

            doc.add_paragraph("Variant 1 (ToвЂgвЂri): " + correct_answer)
            for i, ans in enumerate(incorrect_answers, 2):
                doc.add_paragraph(f"Variant {i}: {ans}")
            doc.add_paragraph("=" * 50)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            content=buffer
        )
        response['Content-Disposition'] = f'attachment; filename="{test.name}_questions.docx"'
        return response
