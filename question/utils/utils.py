import csv
import json
import random
import tempfile
from pathlib import Path
from io import BytesIO
import base64
import shutil
import subprocess
import mimetypes
import html as std_html
from urllib.parse import unquote
from django.contrib.auth.decorators import login_required
from django.core.files.base import ContentFile
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404
from django.utils.decorators import method_decorator
from django.views import View
from docx.shared import Inches
from openpyxl import load_workbook, Workbook
import re
from docx import Document
from sympy import sympify, latex
import matplotlib.pyplot as plt
from lxml import html as lxml_html


def _word_doc_to_docx_via_com(doc_path):
    try:
        import win32com.client
    except Exception as e:
        raise RuntimeError("MS Word (win32com) topilmadi!") from e

    from pathlib import Path
    doc_path = Path(doc_path).resolve()
    out_path = doc_path.with_suffix(".docx")

    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(doc_path))
        doc.SaveAs(str(out_path), FileFormat=16)
    finally:
        try:
            if doc: doc.Close(False)
        except: pass
        try:
            if word: word.Quit()
        except: pass
    return out_path


def _word_doc_to_docx(doc_path):
    import shutil
    import subprocess
    from pathlib import Path
    import platform

    if platform.system() == "Windows":
        try:
            return _word_doc_to_docx_via_com(doc_path)
        except Exception:
            pass

    soffice = shutil.which("soffice")
    if not soffice:
        raise RuntimeError("LibreOffice o'rnatilmagan")

    doc_path = Path(doc_path).resolve()
    out_dir = doc_path.parent
    proc = subprocess.run(
        [soffice, "--headless", "--convert-to", "docx", "--outdir", str(out_dir), str(doc_path)],
        capture_output=True, text=True, check=False
    )
    out_path = doc_path.with_suffix(".docx")
    if not out_path.exists():
        raise RuntimeError("LibreOffice konvertatsiyada xatolik berdi")
    return out_path


def _detect_charset(raw):
    import re
    head = raw[:5000].decode("ascii", errors="ignore")
    m = re.search(r"charset=([\w-]+)", head, flags=re.IGNORECASE)
    if m: return m.group(1).strip().strip("\"'")
    return "utf-8"


def _embed_word_html_images(root, html_dir, images_dir):
    import base64
    import mimetypes
    from pathlib import Path
    from urllib.parse import unquote

    def resolve_src(src):
        src = (src or "").strip()
        if not src or src.lower().startswith("data:"): return None
        src = unquote(src)
        if src.lower().startswith("file:///"): return Path(src[8:])
        if src.lower().startswith("file:"): return Path(src[5:])
        src_norm = src.replace("\\", "/")
        try:
            for c in [html_dir / src_norm, images_dir / src_norm, images_dir / Path(src_norm).name, html_dir / Path(src_norm).name]:
                if c.exists(): return c
        except Exception:
            pass
        return None

    def embed_attr(el, attr):
        src = el.get(attr) or ""
        fp = resolve_src(src)
        if fp and fp.exists():
            mime = mimetypes.guess_type(str(fp))[0] or "application/octet-stream"
            b64 = base64.b64encode(fp.read_bytes()).decode("ascii")
            el.set(attr, f"data:{mime};base64,{b64}")

    for img in root.xpath("//*[local-name()='img']"):
        embed_attr(img, "src")
    for vml in root.xpath("//*[local-name()='imagedata']"):
        if vml.get("src"): embed_attr(vml, "src")


def _docx_to_html_embedded_via_word(docx_path):
    try:
        import win32com.client
    except Exception as e:
        raise RuntimeError("win32com yo'q") from e

    import tempfile
    from pathlib import Path
    from lxml import html as lxml_html
    docx_path = Path(docx_path).resolve()

    with tempfile.TemporaryDirectory(prefix="quiz_word_html_") as td:
        td_path = Path(td)
        out_html = td_path / "input.html"

        word = None
        doc = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(str(docx_path))
            doc.SaveAs(str(out_html), FileFormat=10)
        finally:
            try:
                if doc: doc.Close(False)
            except: pass
            try:
                if word: word.Quit()
            except: pass

        if not out_html.exists():
            raise RuntimeError("Word HTML yaratmadi.")

        raw = out_html.read_bytes()
        html_text = raw.decode(_detect_charset(raw), errors="replace")
        root = lxml_html.fromstring(html_text)
        images_dir = out_html.parent / f"{out_html.stem}.files"
        _embed_word_html_images(root, out_html.parent, images_dir)
        return lxml_html.tostring(root, encoding="unicode", method="html")


def _docx_to_html_embedded_via_libreoffice(docx_path):
    import subprocess
    import shutil
    import tempfile
    from pathlib import Path
    from lxml import html as lxml_html

    soffice = shutil.which("soffice")
    if not soffice:
        raise RuntimeError("LibreOffice (soffice) topilmadi!")

    docx_path = Path(docx_path).resolve()
    with tempfile.TemporaryDirectory(prefix="quiz_word_html_") as td:
        subprocess.run([soffice, "--headless", "--convert-to", "html", "--outdir", td, str(docx_path)], capture_output=True)
        out_html = Path(td) / f"{docx_path.stem}.html"
        if not out_html.exists():
            out_html = Path(td) / f"{docx_path.stem}.htm"
        if not out_html.exists():
            raise RuntimeError("LibreOffice HTML formatda muvaffaqiyatli saqlay olmadi.")

        raw = out_html.read_bytes()
        html_text = raw.decode(_detect_charset(raw), errors="replace")
        root = lxml_html.fromstring(html_text)
        _embed_word_html_images(root, out_html.parent, out_html.parent)
        return lxml_html.tostring(root, encoding="unicode", method="html")


def _docx_to_html_embedded(docx_path):
    import platform
    if platform.system() == "Windows":
        try:
            return _docx_to_html_embedded_via_word(docx_path)
        except Exception as e:
            print(f"Word orqali HTML qilish o'xshamadi: {e}, LibreOffice usuliga o'tamiz")
            
    try:
        return _docx_to_html_embedded_via_libreoffice(docx_path)
    except Exception as e:
        raise RuntimeError("HTML rasm va formulalar eksporti Windowsda Word, Linuxda LibreOffice orqali ishlashi sozlangan, ammo tizimda bularni topilmadi.")


def parse_word(file) -> list[dict]:
    import tempfile
    import re
    from pathlib import Path
    from lxml import html as lxml_html

    name = getattr(file, "name", "") or ""
    suffix = Path(name).suffix.lower()
    if suffix not in [".docx", ".doc"]:
        raise ValueError("Faqat .docx yoki .doc fayllar qabul qilinadi")

    with tempfile.TemporaryDirectory(prefix="quiz_word_") as td:
        img_path = Path(td)
        in_path = img_path / f"input{suffix}"
        
        with open(in_path, "wb") as f:
            if hasattr(file, "chunks"):
                for chunk in file.chunks(): f.write(chunk)
            else:
                f.write(file.read())

        docx_path = in_path
        if suffix == ".doc":
            docx_path = _word_doc_to_docx(in_path)

        html_text = _docx_to_html_embedded(docx_path)
        root = lxml_html.fromstring(html_text)
        body_list = root.xpath("//body")
        if not body_list:
            return []
        body = body_list[0]

        q_re = re.compile(r"^Savol\s*(\d+)\s*[:.)]?\s*", re.IGNORECASE)
        v_re = re.compile(r"^Variant\s*(\d+)\b", re.IGNORECASE)
        sep_re = re.compile(r"^[=\-]{10,}$")

        questions = []
        current_q_text = []
        current_answers = []

        def iter_blocks(el):
            for child in el.iterchildren():
                tag = str(child.tag).lower()
                if "}" in tag: tag = tag.split("}", 1)[1]
                if tag in ("p", "h1", "h2", "h3", "h4", "table", "div"):
                    yield child
                elif tag in ("ul", "ol"):
                    for li in child.xpath(".//*[local-name()='li']"):
                        yield li
                else:
                    yield from iter_blocks(child)

        for block in iter_blocks(body):
            text = " ".join((block.text_content() or "").split())
            has_img = bool(block.xpath(".//*[local-name()='img']") or block.xpath(".//*[local-name()='math']") or block.xpath(".//*[local-name()='imagedata']"))
            if not text and not has_img:
                continue

            block_html = lxml_html.tostring(block, encoding="unicode", method="html")

            # Separator
            if text == "+++++" or sep_re.match(text):
                if current_q_text and current_answers:
                    questions.append({
                        "text": "\n".join(current_q_text),
                        "image_base64": None,
                        "answers": current_answers
                    })
                current_q_text = []
                current_answers = []
                continue

            # Variant (Eski)
            vm = v_re.match(text)
            if vm:
                vnum = int(vm.group(1))
                t_low = text.lower()
                is_correct = (vnum == 1) or ("to'g'ri" in t_low) or ("togri" in t_low)
                current_answers.append((block_html, is_correct))
                continue

            # Yangi Format
            if text.startswith("====="):
                ans_text = text[5:].strip()
                is_correct = False
                if ans_text.startswith("#"):
                    is_correct = True
                    
                clean_html = block_html.replace("=====", "", 1)
                if is_correct:
                    clean_html = clean_html.replace("#", "", 1)
                    
                current_answers.append((clean_html, is_correct))
                continue

            # Savol bormi
            if q_re.match(text):
                if current_q_text and current_answers:
                    questions.append({
                        "text": "\n".join(current_q_text),
                        "image_base64": None,
                        "answers": current_answers
                    })
                current_q_text = [block_html]
                current_answers = []
                continue

            current_q_text.append(block_html)

        if current_q_text and current_answers:
            questions.append({
                "text": "\n".join(current_q_text),
                "image_base64": None,
                "answers": current_answers
            })

        if questions:
            return questions

        # Fallback jadvali
        tables = body.xpath(".//*[local-name()='table']")
        out2 = []
        for table in tables:
            for tr in table.xpath(".//tr"):
                cells = tr.xpath("./th|./td")
                if len(cells) < 3:
                    continue
                
                stem_cell = cells[0]
                option_cells = cells[1:5]
                
                def cell_to_html(cell):
                    parts = []
                    for child in cell:
                        parts.append(lxml_html.tostring(child, encoding="unicode", method="html"))
                    inner = "".join(parts).strip()
                    if not inner:
                        txt = " ".join((cell.text_content() or "").split())
                        if txt: inner = f"<p>{txt}</p>"
                    return inner if inner else ""
                
                q_html = cell_to_html(stem_cell)
                if not q_html: continue
                
                answers = []
                last_txt = None
                for c in option_cells:
                    c_html = cell_to_html(c)
                    if c_html and c_html != last_txt:
                        answers.append((c_html, False))
                        last_txt = c_html
                        
                if len(answers) < 2:
                    continue
                    
                answers[0] = (answers[0][0], True)
                out2.append({
                    "text": q_html,
                    "image_base64": None,
                    "answers": answers
                })

        return out2

from question.models import Test, Question, Answer


def parse_excel(file, additional_files=None):
    wb = load_workbook(file)
    sheet = wb.active
    questions = []

    def convert_formula_to_base64(formula):
        """LaTeX yoki matematik formulani rasmga aylantirib, Base64 ga kodlaydi."""
        try:
            # Formulani LaTeX formatiga oвЂtkazish
            expr = sympify(formula)
            latex_expr = latex(expr)

            # Rasm yaratish
            plt.figure(figsize=(4, 1))
            plt.text(0.5, 0.5, f"${latex_expr}$", fontsize=12, ha='center', va='center')
            plt.axis('off')

            # Rasmni Base64 ga aylantirish
            buffer = BytesIO()
            plt.savefig(buffer, format='png', bbox_inches='tight', dpi=100)
            plt.close()
            buffer.seek(0)
            image_base64 = base64.b64encode(buffer.read()).decode('utf-8')
            return image_base64
        except Exception as e:
            print(f"Formula konvertatsiyasida xato: {str(e)}")
            return None

    for row in sheet.iter_rows(min_row=2, values_only=True):
        # Qator boвЂsh yoki yetarli ustunlarga ega emasligini tekshirish
        if not row or len(row) < 2:  # Kamida savol + 1 javob boвЂlishi kerak
            print(f"Qator oвЂtkazib yuborildi: yetarli ustunlar yoвЂq (ustunlar soni: {len(row) if row else 0})")
            continue

        question_text = row[0]
        if not question_text:
            print("Qator oвЂtkazib yuborildi: savol matni boвЂsh")
            continue

        # Formula aniqlash (masalan, LaTeX yoki oddiy matematik ifoda)
        formula_match = re.search(r'\[Formula:\s*([^\]]+)\]', question_text)
        image_base64 = None
        if formula_match:
            formula = formula_match.group(1).strip()
            image_base64 = convert_formula_to_base64(formula)
            # Formula belgisini matndan olib tashlash
            question_text = re.sub(r'\[Formula:\s*[^\]]+\]', '', question_text).strip()

        # [Rasm: nom.jpg] belgisini olib tashlash (agar mavjud boвЂlsa)
        question_text = re.sub(r'\[Rasm:\s*[^\]]+\]', '', question_text).strip()

        # Javoblarni olish (dinamik son)
        answers = [(row[i], i == 1) for i in range(1, len(row)) if row[i]]  # 1-variant toвЂgвЂri
        if len(answers) < 2:  # Kamida 1 toвЂgвЂri va 1 notoвЂgвЂri javob
            print(f"Qator oвЂtkazib yuborildi: yetarli javoblar yoвЂq (javoblar soni: {len(answers)})")
            continue

        # Javoblarni tasodifiy tartibda almashtirish
        random.shuffle(answers)

        questions.append({
            'text': question_text,
            'image_base64': image_base64,
            'answers': answers
        })

    print(f"Jami {len(questions)} ta savol topildi")
    return questions


@method_decorator(login_required, name='dispatch')
class ImportQuestionsView(View):
    def post(self, request, test_id):
        print(f"POST soвЂrovi keldi: test_id={test_id}, foydalanuvchi={request.user}")
        test = get_object_or_404(Test, id=test_id)
        print(f"Test topildi: {test.name} (ID: {test_id})")

        file = request.FILES.get('import_file')
        additional_files = request.FILES.getlist('additional_files', [])
        print(f"Fayl: {file.name if file else 'Fayl yuklanmadi'}")
        print(f"QoвЂshimcha fayllar: {[f.name for f in additional_files] if additional_files else 'YoвЂq'}")

        if not file:
            return JsonResponse({"success": False, "errors": "Fayl yuklanmadi"})

        name = (file.name or "").lower()
        is_excel = name.endswith(".xlsx")
        is_word = name.endswith(".docx") or name.endswith(".doc")
        if not (is_excel or is_word):
            return JsonResponse({"success": False, "errors": "Faqat .xlsx, .docx yoki .doc fayllar qabul qilinadi"})

        if file.size > 10 * 1024 * 1024:
            print(f"Xato: Fayl hajmi {file.size} bayt, 10MB dan katta")
            return JsonResponse({"success": False, "errors": "Fayl hajmi 10MB dan kichik boвЂlishi kerak"})

        try:
            if is_excel:
                print("Excel faylini parsing boshlandi...")
                questions = parse_excel(file, additional_files)
            else:
                print("Word faylini parsing boshlandi...")
                questions = parse_word(file)
            print(f"Parsing natijasi: {len(questions)} ta savol topildi")
            if not questions:
                print("Xato: Faylda savollar topilmadi")
                return JsonResponse({"success": False, "errors": "Faylda savollar topilmadi"})

            # Savollarni sessionвЂ™da saqlash
            request.session['imported_questions'] = questions
            request.session['test_id'] = test_id
            print(f"SessionвЂ™da saqlandi: {len(questions)} ta savol, test_id={test_id}")

            return JsonResponse({
                "success": True,
                "questions": questions,
                "message": "Savollar muvaffaqiyatli oвЂqildi, modalda koвЂrsatilmoqda!"
            })
        except Exception as e:
            print(f"Xato yuz berdi: {str(e)}")
            return JsonResponse({"success": False, "errors": f"Faylni oвЂqishda xato: {str(e)}"})


@method_decorator(login_required, name='dispatch')
class DownloadTemplateView(View):
    def get(self, request):
        wb = Workbook()
        ws = wb.active
        ws.title = "Questions Template"

        headers = ["Savol matni", "Variant 1", "Variant 2", "Variant 3", "Variant 4", "ToвЂgвЂri javob"]
        ws.append(headers)

        example_row = ["Misol savol", "ToвЂgвЂri javob", "NotoвЂgвЂri 1", "NotoвЂgвЂri 2", "NotoвЂgвЂri 3", "1 variant"]
        ws.append(example_row)

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            content=buffer,
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename=questions_template.xlsx'
        return response


@method_decorator(login_required, name='dispatch')
class SaveImportedQuestionsView(View):
    def post(self, request):
        test_id = request.session.get('test_id')
        if not test_id:
            return JsonResponse({"success": False, "errors": "Test topilmadi"})

        questions = request.session.get('imported_questions', [])
        if not questions:
            return JsonResponse({"success": False, "errors": "Saqlash uchun savollar topilmadi"})

        test = get_object_or_404(Test, id=test_id)
        saved_count = 0

        try:
            for q in questions:
                question_text = q['text'].strip()
                correct_answer = next((text for text, is_correct in q['answers'] if is_correct), None)
                if not correct_answer:
                    print(f"Savol oвЂtkazib yuborildi: toвЂgвЂri javob yoвЂq: {question_text}")
                    continue

                # Bazada takrorlanishni tekshirish
                exists = Question.objects.filter(
                    test=test,
                    text__iexact=question_text
                ).filter(
                    answers__text__iexact=correct_answer,
                    answers__is_correct=True
                ).exists()

                if exists:
                    print(f"Savol oвЂtkazib yuborildi: takrorlangan savol: {question_text} (toвЂgвЂri javob: {correct_answer})")
                    continue

                # Yangi savolni saqlash
                question = Question.objects.create(test=test, text=question_text)
                if q.get('image_base64'):
                    image_data = base64.b64decode(q['image_base64'])
                    image_name = f"question_{question.id}.png"
                    question.image.save(image_name, ContentFile(image_data))

                for answer_text, is_correct in q['answers']:
                    Answer.objects.create(
                        question=question,
                        text=answer_text,
                        is_correct=is_correct
                    )

                saved_count += 1

            # SessionвЂ™ni tozalash
            del request.session['imported_questions']
            del request.session['test_id']
            return JsonResponse({
                "success": True,
                "message": f"{saved_count} ta savol muvaffaqiyatli saqlandi!"
            })
        except Exception as e:
            print(f"Saqlashda xato: {str(e)}")
            return JsonResponse({"success": False, "errors": f"Saqlashda xato: {str(e)}"})


@method_decorator(login_required, name='dispatch')
class DownloadTemplateView(View):
    def get(self, request):
        wb = Workbook()
        ws = wb.active
        ws.title = "Questions Template"

        headers = ["Savol matni", "Variant 1", "Variant 2", "Variant 3", "Variant 4"]  # ToвЂgвЂri javob ustuni olib tashlandi
        ws.append(headers)

        example_row = ["Misol savol", "ToвЂgвЂri javob", "NotoвЂgвЂri 1", "NotoвЂgвЂri 2", "NotoвЂgвЂri 3"]
        ws.append(example_row)

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            content=buffer,
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename=questions_template.xlsx'
        return response


@method_decorator(login_required, name='dispatch')
class ExportQuestionsView(View):
    def get(self, request, test_id, format_type):
        test = get_object_or_404(Test, id=test_id)
        questions = Question.objects.filter(test=test).prefetch_related('answers')

        if format_type == 'json':
            return self.export_json(questions, test)
        elif format_type == 'xlsx':
            return self.export_excel(questions, test)
        elif format_type == 'csv':
            return self.export_csv(questions, test)
        elif format_type == 'docx':
            return self.export_docx(questions, test)
        else:
            return HttpResponse(status=400, content="NotoвЂgвЂri format turi")

    def export_json(self, questions, test):
        data = {
            'test_name': test.name,
            'questions': [
                {
                    'text': q.text,
                    'image': q.image.url if q.image else None,
                    'answers': [
                        {'text': a.text, 'is_correct': a.is_correct}
                        for a in q.answers.all()
                    ]
                }
                for q in questions
            ]
        }
        response = HttpResponse(
            content_type='application/json',
            content=json.dumps(data, ensure_ascii=False, indent=2)
        )
        response['Content-Disposition'] = f'attachment; filename="{test.name}_questions.json"'
        return response

    def export_excel(self, questions, test):
        wb = Workbook()
        ws = wb.active
        ws.title = "Questions"

        headers = ["Savol matni", "Variant 1", "Variant 2", "Variant 3", "Variant 4", "ToвЂgвЂri javob", "Rasm"]
        ws.append(headers)

        for q in questions:
            answers = list(q.answers.all())
            correct_answer = next((a.text for a in answers if a.is_correct), '')
            incorrect_answers = [a.text for a in answers if not a.is_correct][:3]
            row = [
                q.text,
                correct_answer,  # Variant 1 doim toвЂgвЂri
                incorrect_answers[0] if len(incorrect_answers) > 0 else '',
                incorrect_answers[1] if len(incorrect_answers) > 1 else '',
                incorrect_answers[2] if len(incorrect_answers) > 2 else '',
                "1 variant",
                q.image.name if q.image else ''
            ]
            ws.append(row)

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            content=buffer
        )
        response['Content-Disposition'] = f'attachment; filename="{test.name}_questions.xlsx"'
        return response

    def export_csv(self, questions, test):
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{test.name}_questions.csv"'

        writer = csv.writer(response, lineterminator='\n')
        writer.writerow(["Savol matni", "Variant 1", "Variant 2", "Variant 3", "Variant 4", "ToвЂgвЂri javob", "Rasm"])

        for q in questions:
            answers = list(q.answers.all())
            correct_answer = next((a.text for a in answers if a.is_correct), '')
            incorrect_answers = [a.text for a in answers if not a.is_correct][:3]
            row = [
                q.text,
                correct_answer,
                incorrect_answers[0] if len(incorrect_answers) > 0 else '',
                incorrect_answers[1] if len(incorrect_answers) > 1 else '',
                incorrect_answers[2] if len(incorrect_answers) > 2 else '',
                "1 variant",
                q.image.name if q.image else ''
            ]
            writer.writerow(row)

        return response

    def export_docx(self, questions, test):
        doc = Document()
        doc.add_heading(f"{test.name} - Savollar", 0)

        for idx, q in enumerate(questions, 1):
            doc.add_paragraph(f"Savol {idx}: {q.text}")
            if q.image:
                try:
                    with q.image.open('rb') as img_file:
                        image_data = base64.b64encode(img_file.read()).decode('utf-8')
                    buffer = BytesIO(base64.b64decode(image_data))
                    doc.add_picture(buffer, width=Inches(4))
                except Exception:
                    doc.add_paragraph(f"[Rasm: {q.image.name}]")

            answers = list(q.answers.all())
            correct_answer = next((a.text for a in answers if a.is_correct), '')
            incorrect_answers = [a.text for a in answers if not a.is_correct][:3]

            doc.add_paragraph("Variant 1 (ToвЂgвЂri): " + correct_answer)
            for i, ans in enumerate(incorrect_answers, 2):
                doc.add_paragraph(f"Variant {i}: {ans}")
            doc.add_paragraph("=" * 50)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            content=buffer
        )
        response['Content-Disposition'] = f'attachment; filename="{test.name}_questions.docx"'
        return response
